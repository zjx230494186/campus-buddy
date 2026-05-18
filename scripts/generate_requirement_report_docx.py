from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = r"D:\big_homework\deliverables\需求分析报告-校园搭子.docx"


def set_east_asia(run, name="微软雅黑"):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    set_east_asia(run)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def paragraph(doc, text="", *, size=12, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
              first_line=True, space_after=6, space_before=0, line_spacing=1.45):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    set_east_asia(run)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18 if level == 1 else 14 if level == 2 else 12)
    set_east_asia(run)
    run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    run.font.size = Pt(11)
    set_east_asia(run)
    return p


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].width = Cm(widths[i])
        set_cell_text(hdr[i], h, bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr[i], "D9E2F3")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].width = Cm(widths[i])
            set_cell_text(cells[i], v, size=10.2)
    doc.add_paragraph()
    return t


def figure_placeholder(doc, fig_no, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"【插图占位：图 {fig_no}  {title}】")
    run.bold = True
    run.font.size = Pt(10.5)
    set_east_asia(run)
    run.font.color.rgb = RGBColor(79, 98, 40)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    run2 = p2.add_run(f"图 {fig_no}  {title}")
    run2.font.size = Pt(10)
    set_east_asia(run2)
    run2.font.color.rgb = RGBColor(96, 96, 96)


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(76)
    r = p.add_run("校园搭子平台\n需求分析报告")
    r.bold = True
    r.font.size = Pt(24)
    set_east_asia(r)
    r.font.color.rgb = RGBColor(31, 78, 121)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(18)
    r2 = p2.add_run("软件系统分析与设计课程作业")
    r2.font.size = Pt(14)
    set_east_asia(r2)

    rows = [
        ("项目名称", "校园搭子"),
        ("文档阶段", "需求分析"),
        ("编写人", "[待填写]"),
        ("小组成员", "[待填写]"),
        ("指导教师", "[待填写]"),
        ("日期", "2026 年 5 月"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.autofit = False
    for k, v in rows:
        cells = t.add_row().cells
        cells[0].width = Cm(4.2)
        cells[1].width = Cm(10.8)
        set_cell_text(cells[0], k, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cells[0], "D9E2F3")
        set_cell_text(cells[1], v)
    doc.add_page_break()


def toc(doc):
    heading(doc, "目录", 1)
    items = [
        "1 项目概述",
        "2 业务背景与问题定义",
        "3 利益相关者与用户分析",
        "4 业务范围与产品边界",
        "5 业务流程与用例分析",
        "6 需求分层与功能需求",
        "7 数据对象与规则约束",
        "8 非功能需求",
        "9 版本范围与实施基线",
        "10 附录",
    ]
    for item in items:
        bullet(doc, item)
    doc.add_page_break()


def configure(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(12)

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("校园搭子平台需求分析报告")
    run.font.size = Pt(9)
    set_east_asia(run)
    run.font.color.rgb = RGBColor(130, 130, 130)


def build():
    doc = Document()
    configure(doc)
    cover(doc)
    toc(doc)

    heading(doc, "1 项目概述", 1)
    paragraph(doc, "校园搭子平台面向高校学生在校内寻找结伴对象的实际需求，统一承接就餐、学习、运动、课程作业组队和大创项目组队五类核心场景。系统通过结构化需求发布、广场发现、规则推荐、低压力联系、信用评价与基础治理机制，建立一条从需求产生到合作完成的稳定闭环。")
    paragraph(doc, "本版本按单学校、单平台、网页端产品形态设计，优先保证从账号准入、需求发布、审核展示、联系确认到评价沉淀的主链路清晰可实施，为后续详细设计、建模落地和最小闭环实现提供统一基线。")

    heading(doc, "1.1 项目目标", 2)
    table(doc, ["目标类别", "目标内容"], [
        ("效率目标", "替代群聊刷屏和熟人转介绍方式，提升学生在校内寻找搭子的匹配效率。"),
        ("匹配目标", "以时间、地点、标签、技能、水平等字段为核心，提高需求发现和对象筛选准确性。"),
        ("风险控制目标", "通过身份校验、先审后发、信用展示和治理流程降低骚扰、失联、爽约和虚假信息风险。"),
        ("实施目标", "形成可直接进入详细设计和编码拆解的需求基线，保证需求可追踪、可建模、可实现。"),
    ], [4.0, 11.2])

    heading(doc, "1.2 产品定位", 2)
    paragraph(doc, "产品定位为校内轻量级搭子匹配平台，而非泛社交社区。平台聚焦任务型和场景型连接，不承担复杂社区关系链、内容运营、支付担保和重型即时通信职能。设计重点是让用户能够在短时间内找到合适对象，并在合作前后形成可积累的信任信息。")

    heading(doc, "2 业务背景与问题定义", 1)
    paragraph(doc, "高校学生在日常校园生活中存在持续、分散而高频的结伴需求，但现有承载方式主要依赖班群、宿舍群、朋友圈和熟人介绍。这类方式在信息触达、条件筛选、响应速度和风险识别方面均存在显著不足。")
    table(doc, ["问题编号", "问题定义", "影响"], [
        ("RR1", "现有找搭子渠道分散且触达效率低", "有效需求被聊天流淹没，错过时机，发布成本高。"),
        ("RR2", "缺少按条件筛选合适对象的能力", "即使找到人，也容易出现时间、地点、能力和目标不匹配。"),
        ("RR3", "陌生人建立联系存在心理门槛", "用户不愿过早暴露联系方式，意向建立过程不稳定。"),
        ("RR4", "缺少对爽约、划水和失联的识别与约束", "合作质量不可控，重复试错成本高。"),
        ("RR5", "历史合作关系无法沉淀复用", "每次寻找对象都需要重新筛选，难以形成稳定协作网络。"),
        ("RR6", "平台缺少对违规和纠纷的治理抓手", "骚扰、虚假信息、恶意评价和申诉纠偏无法形成闭环。"),
    ], [1.7, 7.0, 6.5])

    heading(doc, "3 利益相关者与用户分析", 1)
    heading(doc, "3.1 利益相关者", 2)
    table(doc, ["角色", "核心诉求", "系统价值"], [
        ("学生用户", "快速找到合适搭子、降低沟通成本、识别风险", "平台主要服务对象，决定需求密度和活跃度"),
        ("被联系对象", "筛选请求、保护隐私、判断对方可靠性", "决定联系链路和转化效率"),
        ("管理员", "审核内容、处理举报、处理申诉、治理账号", "维持平台秩序与规则执行"),
        ("项目团队", "统一需求基线、降低返工、形成可实施方案", "保障设计与实现一致性"),
    ], [3.0, 6.2, 6.0])

    heading(doc, "3.2 目标用户画像", 2)
    paragraph(doc, "目标用户以在校本科生和研究生为主，需求具有明显场景化和时效性。按照需求稳定性和合作周期，可划分为临时结伴型用户和协作组队型用户两类。前者更关注快速响应与便利性，后者更关注匹配质量、信用约束和长期协同。")
    table(doc, ["画像", "典型场景", "核心诉求", "主要风险点"], [
        ("P1 临时就餐型", "饭点寻找同食堂饭搭子", "同时间、同地点、快速响应", "需求触达慢、消息被淹没"),
        ("P2 学习陪伴型", "备考、自习、监督学习", "目标一致、作息相近、可持续", "作息不匹配、长期稳定性差"),
        ("P3 运动结伴型", "羽毛球、跑步等运动搭子", "水平接近、地点合适、爽约少", "临时放鸽子、水平差异过大"),
        ("P4 协作组队型", "课程作业或大创项目组队", "技能匹配、责任心、可交付", "划水、拖延、能力失真"),
    ], [3.0, 4.4, 4.6, 4.2])
    figure_placeholder(doc, "3-1", "用户画像图")

    heading(doc, "4 业务范围与产品边界", 1)
    heading(doc, "4.1 核心业务范围", 2)
    bullets = [
        "账号准入与身份校验：校园邮箱验证码登录、基础身份认证。",
        "结构化需求发布：按统一字段模板发布五类搭子需求。",
        "内容审核：需求先进入审核，再进入公开广场。",
        "广场发现：浏览、筛选、搜索、排序和规则推荐。",
        "联系链路：邀约、双向确认、站内消息和联系方式渐进解锁。",
        "评价沉淀：互评、信用摘要、历史合作记录。",
        "治理闭环：举报、申诉、被投诉方回应、管理员处理。",
    ]
    for b in bullets:
        bullet(doc, b)

    heading(doc, "4.2 非本阶段范围", 2)
    bullets = [
        "复杂算法推荐模型和个性化训练体系。",
        "完整即时通讯系统，包括多端同步、已读回执和实时到达保障。",
        "支付、担保、交易撮合和线下履约管理。",
        "多学校统一运营能力和跨校匹配。",
        "内容社区、动态流和泛社交关系链扩展。",
    ]
    for b in bullets:
        bullet(doc, b)

    heading(doc, "4.3 核心场景定义", 2)
    table(doc, ["场景", "发布目的", "关键字段"], [
        ("饭搭子", "就餐结伴", "时间段、食堂、预算、口味偏好"),
        ("学习搭子", "自习、备考、监督", "目标、作息、频率、监督方式"),
        ("运动搭子", "运动结伴", "项目、水平、场地、时间"),
        ("小组作业队友", "课程作业组队", "课程、角色、技能、交付周期"),
        ("大创项目队友", "项目协作", "方向、角色、技能、投入周期"),
    ], [3.0, 4.4, 8.0])

    heading(doc, "5 业务流程与用例分析", 1)
    heading(doc, "5.1 业务主流程", 2)
    paragraph(doc, "产品主流程分为学生侧业务主流程和管理侧治理主流程。学生侧围绕“准入 - 发布 - 展示 - 联系 - 合作 - 评价 - 沉淀”展开；管理侧围绕“审核 - 举报 - 申诉 - 处置”展开。")
    figure_placeholder(doc, "5-1", "业务总流程图")

    heading(doc, "5.2 学生侧流程", 2)
    table(doc, ["阶段", "输入", "输出"], [
        ("账号准入", "校园邮箱、验证码、基础身份信息", "可用账号、认证状态"),
        ("需求发布", "场景类型、时间地点、标签、描述", "待审核需求"),
        ("广场发现", "搜索条件、筛选条件、推荐结果", "候选需求或候选对象"),
        ("联系确认", "邀约、留言、站内消息", "已确认意向或建立会话"),
        ("合作完成", "线下活动或线上协作结果", "待评价关系"),
        ("评价沉淀", "互评内容、信用摘要、历史记录", "可供后续筛选的信任信息"),
    ], [3.0, 6.1, 5.1])

    heading(doc, "5.3 管理侧流程", 2)
    table(doc, ["阶段", "输入", "处理动作", "输出"], [
        ("内容审核", "待审核需求", "通过、驳回、备注", "公开需求或驳回结果"),
        ("举报处理", "用户举报、消息举报、需求举报", "核验、警告、限制、冻结", "案件结果"),
        ("申诉处理", "申诉说明、证据、被投诉方回应", "复核、维持、纠正", "最终裁定"),
    ], [2.5, 4.4, 4.4, 3.0])

    heading(doc, "5.4 用例结构", 2)
    paragraph(doc, "需求分析阶段将用例图拆分为一个总体用例图和三个业务模块图。总体图用于描述角色与主流程关系；模块图分别覆盖需求发布与匹配、联系确认、信用沉淀与关系复用。")
    figure_placeholder(doc, "5-2", "总体用例图")
    figure_placeholder(doc, "5-3", "模块用例图一：需求发布与匹配")
    figure_placeholder(doc, "5-4", "模块用例图二：邀约与双向确认")
    figure_placeholder(doc, "5-5", "模块用例图三：信用评价与历史记录")

    heading(doc, "6 需求分层与功能需求", 1)
    paragraph(doc, "需求结构按原始需求、系统特性、研发需求和用户故事四层组织。四层之间保持一一追踪关系，其中原始需求定义问题域，系统特性定义能力域，研发需求定义实现边界，用户故事定义用户视角的可交付目标。")
    figure_placeholder(doc, "6-1", "需求树总览图")

    heading(doc, "6.1 系统特性", 2)
    table(doc, ["系统特性", "定义", "承接问题"], [
        ("SF1 需求发布系统", "承接结构化发布、提交和进入审核队列", "RR1、RR2"),
        ("SF2 智能匹配与推荐系统", "承接筛选、排序、推荐和广场发现", "RR1、RR2"),
        ("SF3 邀约与双向确认系统", "承接低压力联系、邀约、站内消息和联系确认", "RR3"),
        ("SF4 信用评价系统", "承接互评、信用摘要、推荐权重影响和申诉纠偏", "RR4、RR6"),
        ("SF5 搭子记录与关系管理系统", "承接历史记录、关系沉淀、再次联系和治理留痕", "RR5、RR6"),
    ], [4.0, 7.0, 4.2])
    figure_placeholder(doc, "6-2", "系统特性截图")

    heading(doc, "6.2 研发需求", 2)
    table(doc, ["编号", "研发需求", "所属系统特性"], [
        ("IR1", "用户可以创建搭子需求，填写标题和详细描述", "SF1"),
        ("IR2", "发布需求时支持选择学习、运动、吃饭、项目组队、出行、娱乐等场景标签", "SF1"),
        ("IR3", "发布需求时支持设置时间信息和紧急程度", "SF1"),
        ("IR4", "发布需求时支持设置地点和人数上限", "SF1"),
        ("IR5", "系统根据用户需求标签自动匹配条件相近的其他用户", "SF2"),
        ("IR6", "匹配结果可按时间重合度排序展示", "SF2"),
        ("IR7", "匹配结果可按距离远近排序展示", "SF2"),
        ("IR8", "用户可向匹配对象发起组队意向并附带简短留言", "SF3"),
        ("IR9", "接收方看到邀请时对方信息匿名展示，接受后显示联系方式", "SF3"),
        ("IR10", "支持一键拒绝并附带默认话术，拒绝后双方互不可见", "SF3"),
        ("IR11", "搭子活动结束后双方互评，评价维度包含准时性和合作态度", "SF4"),
        ("IR12", "用户个人主页展示信用分和近期评价摘要", "SF4"),
        ("IR13", "低信用分用户在匹配推荐中降低权重", "SF4"),
        ("IR14", "用户可以查看历史搭子记录列表", "SF5"),
        ("IR15", "管理员可查看和处理用户举报", "SF5"),
        ("IR16", "管理员可对违规账号进行冻结处理", "SF5"),
        ("IR17", "用户可通过校园邮箱验证码完成登录或注册", "SF1"),
        ("IR18", "用户需提交学号及基础校内信息完成身份认证", "SF1"),
        ("IR19", "搭子需求提交后需经管理员审核通过方可进入广场", "SF1"),
        ("IR20", "用户可通过站内消息发送文本、图片、截图和联系方式卡片", "SF3"),
        ("IR21", "用户可针对评价、举报处理或纠纷结果提交申诉并上传证据", "SF4"),
        ("IR22", "被投诉方可查看案件并补充说明与证据回应", "SF5"),
    ], [1.6, 10.2, 3.6])
    figure_placeholder(doc, "6-3", "研发需求截图")

    heading(doc, "6.3 用户故事", 2)
    table(doc, ["编号", "用户故事", "优先级", "对应需求"], [
        ("US1", "快速匹配同时间同食堂的饭搭子", "高", "IR1-IR7"),
        ("US2", "发布的需求能够被真正有空的人稳定看到", "高", "IR1、IR19"),
        ("US3", "按学习目标和作息时间匹配固定学伴", "高", "IR1-IR7"),
        ("US4", "提前识别并避开经常爽约的对象", "高", "IR11-IR13"),
        ("US5", "按运动水平和场地筛选运动伙伴", "中", "IR1-IR7"),
        ("US6", "发布运动需求后尽快收到响应", "中", "IR5-IR8"),
        ("US7", "在不暴露外部联系方式的前提下先建立联系意向", "高", "IR8-IR10、IR20"),
        ("US8", "组队时过滤经常划水拖延的人", "高", "IR11-IR13"),
        ("US9", "按技能精准找大创项目队友", "高", "IR1-IR8"),
        ("US10", "通过系统初筛降低沟通成本", "中", "IR5-IR8"),
        ("US11", "收到不想参加的邀请时可以快速拒绝", "中", "IR10"),
        ("US12", "快速找到之前合作愉快的对象再次联系", "中", "IR14"),
        ("US13", "先完成校内身份认证再开始找搭子", "高", "IR17、IR18"),
        ("US14", "需求审核通过后再进入公开广场", "高", "IR19"),
        ("US15", "通过站内消息进一步沟通后再交换联系方式", "高", "IR20"),
        ("US16", "如果被误评或被误处理，可以提交申诉纠正", "高", "IR21、IR22"),
    ], [1.6, 8.4, 1.6, 3.8])
    figure_placeholder(doc, "6-4", "用户故事截图")

    heading(doc, "6.4 关键用户故事展开", 2)
    paragraph(doc, "US2  发布的需求能够被真正有空的人稳定看到", bold=True, first_line=False, space_after=4)
    paragraph(doc, "作为一名需要发布搭子需求的学生，我希望我的需求在审核通过后进入公开广场，从而被真正有相同需求且当前可响应的用户看到。")
    bullet(doc, "验收标准 1：需求提交后状态进入“待审核”，不会直接公开。")
    bullet(doc, "验收标准 2：管理员审核通过后，需求进入广场公开列表并可被浏览。")

    paragraph(doc, "US7  在不暴露外部联系方式的前提下先建立联系意向", bold=True, first_line=False, space_before=6, space_after=4)
    paragraph(doc, "作为一名不希望在建立联系初期暴露个人联系方式的学生，我希望先通过邀约、匿名确认或站内消息表达意向，再决定是否进一步交换联系方式。")
    bullet(doc, "验收标准 1：用户可从需求详情页发起邀约或发送站内消息。")
    bullet(doc, "验收标准 2：在意向未确认前，系统至少提供一种不直接暴露外部联系方式的联系路径。")

    paragraph(doc, "US14  需求审核通过后再进入公开广场", bold=True, first_line=False, space_before=6, space_after=4)
    paragraph(doc, "作为平台用户，我希望公开展示的需求已经过基础审核，从而降低虚假信息、骚扰信息和不规范内容对使用体验的影响。")
    bullet(doc, "验收标准 1：需求状态包含待审核、审核通过、审核驳回。")
    bullet(doc, "验收标准 2：仅审核通过状态的需求可以在广场中公开展示。")

    paragraph(doc, "US16  如果被误评或被误处理，可以提交申诉纠正", bold=True, first_line=False, space_before=6, space_after=4)
    paragraph(doc, "作为被误评、被误举报或被误处理的用户，我希望提交申诉说明并补充证据，以便管理员重新判断案件结果。")
    bullet(doc, "验收标准 1：用户可针对评价或处理结果发起申诉。")
    bullet(doc, "验收标准 2：申诉支持补充说明和上传证据材料。")

    heading(doc, "7 数据对象与规则约束", 1)
    heading(doc, "7.1 核心数据对象", 2)
    table(doc, ["对象", "关键属性", "说明"], [
        ("用户", "账号状态、认证状态、基础资料、信用摘要", "作为发布者、浏览者、联系方和被联系方的统一主体"),
        ("需求", "场景、标题、描述、时间、地点、人数、状态", "承载广场发现和联系入口"),
        ("邀约/会话", "发起方、接收方、状态、消息内容、解锁状态", "承载低压力联系和渐进确认"),
        ("评价", "评价方、被评价方、维度、结果、申诉状态", "承载合作后反馈与信用沉淀"),
        ("举报/申诉案件", "提交方、对象、说明、证据、处理结果", "承载治理与纠偏闭环"),
    ], [2.8, 6.6, 5.8])
    figure_placeholder(doc, "7-1", "核心数据对象关系图")

    heading(doc, "7.2 关键状态定义", 2)
    table(doc, ["对象", "关键状态", "业务含义"], [
        ("用户认证", "未认证、认证中、已认证、认证驳回", "决定发布与联系权限边界"),
        ("需求", "草稿、待审核、已发布、已关闭、驳回", "决定是否进入广场与是否可被联系"),
        ("邀约", "已发起、已接受、已拒绝、已失效", "决定双方是否进入下一步沟通"),
        ("案件", "待处理、处理中、已裁定、已关闭", "决定治理链路进展"),
    ], [2.8, 4.6, 7.8])

    heading(doc, "7.3 业务规则", 2)
    rules = [
        "R1 身份规则：登录以校园邮箱验证码完成，认证以学号和基础校内信息完成。",
        "R2 发布规则：所有公开需求需先审核后展示。",
        "R3 联系规则：联系链路以站内路径优先，外部联系方式交换发生在双方确认之后。",
        "R4 信用规则：互评结果参与信用摘要形成，并影响后续推荐权重。",
        "R5 治理规则：举报、申诉和管理员处理结果需要留痕，保证可追踪和可复核。",
        "R6 场景规则：本阶段以五类核心场景为主；出行、娱乐作为扩展场景保留，不作为本轮实现主线。",
    ]
    for r in rules:
        bullet(doc, r)

    heading(doc, "8 非功能需求", 1)
    table(doc, ["类别", "需求"], [
        ("可用性", "关键流程步骤明确，用户能够在少量操作内完成发布、筛选和联系。"),
        ("一致性", "同一业务对象在报告、建模和后续实现中的命名、状态和字段保持一致。"),
        ("安全性", "用户身份信息、联系方式和案件信息需要具备基础访问控制。"),
        ("可追踪性", "需求、评价、举报和申诉等关键记录应具备查询和复核能力。"),
        ("可扩展性", "保留多学校扩展、更多场景扩展和治理规则扩展空间。"),
    ], [3.0, 12.2])

    heading(doc, "9 版本范围与实施基线", 1)
    paragraph(doc, "当前版本的实施基线以最小闭环为原则，优先保证 M0 与 M1 范围内的业务可落地。实施顺序建议按账号准入、需求发布、审核展示、广场发现、联系确认、信用沉淀、治理闭环逐步展开。")
    table(doc, ["实施优先级", "模块", "目标"], [
        ("P0", "账号准入与身份认证", "建立可用账号体系和发布前置条件"),
        ("P0", "需求发布与审核", "形成待审核到公开展示的稳定路径"),
        ("P0", "广场发现", "让需求被查找、筛选和浏览"),
        ("P1", "邀约与站内联系", "让用户在平台内建立意向和会话"),
        ("P1", "互评与信用摘要", "让合作结果可以沉淀并影响后续匹配"),
        ("P1", "举报与申诉", "形成治理闭环"),
    ], [2.2, 4.8, 8.2])
    paragraph(doc, "该基线可直接作为后续详细设计和任务拆解的输入，其中研发需求层适合映射为接口、状态机与页面功能项，用户故事层适合映射为验收标准与测试用例来源。")

    heading(doc, "10 附录", 1)
    heading(doc, "10.1 插图清单", 2)
    for item in [
        "图 3-1 用户画像图",
        "图 5-1 业务总流程图",
        "图 5-2 总体用例图",
        "图 5-3 模块用例图一：需求发布与匹配",
        "图 5-4 模块用例图二：邀约与双向确认",
        "图 5-5 模块用例图三：信用评价与历史记录",
        "图 6-1 需求树总览图",
        "图 6-2 系统特性截图",
        "图 6-3 研发需求截图",
        "图 6-4 用户故事截图",
        "图 7-1 核心数据对象关系图",
    ]:
        bullet(doc, item)

    heading(doc, "10.2 AI 辅助说明", 2)
    paragraph(doc, "本报告初稿的结构组织、表格整理和文字规范化由生成式人工智能辅助完成，使用工具为 OpenAI Codex。项目边界确认、需求取舍、截图准备、平台录入、文档复核和最终提交由小组成员完成。", first_line=False)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build()
